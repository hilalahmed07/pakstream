#!/usr/bin/env python3
"""
CSO Documentation generator for PakStream.

Generates six formal Word (.docx) deliverables into this folder:
    01_Project_Plan.docx
    02_SRS.docx
    03_Design_Document.docx
    04_Prototype_Document.docx
    05_Use_Case_Document.docx
    06_Test_Cases.docx

Content is derived from the actual PakStream codebase (README.md,
PROJECT_ARCHITECTURE.md, backend models/routes/services, frontend structure).

Usage:
    python3 generate_docs.py

Requires: python-docx  (pip install python-docx)
"""

import os
from datetime import date

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
DEMO_DIR = os.path.normpath(os.path.join(HERE, "..", "..", "demo"))

PROJECT = "PakStream"
SUBTITLE = "Secure Media Streaming & Document Management Platform"
VERSION = "1.0"
DATE_STR = date(2026, 6, 2).strftime("%B %d, %Y")
AUTHOR = "PakStream Project Team"

BRAND = RGBColor(0xE5, 0x09, 0x14)   # Netflix-style red
DARK = RGBColor(0x20, 0x20, 0x20)
GREY = RGBColor(0x66, 0x66, 0x66)


# --------------------------------------------------------------------------- #
# Shared helpers
# --------------------------------------------------------------------------- #
def new_doc(doc_title):
    doc = Document()

    # Base style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Heading colors
    for i in (1, 2, 3):
        h = doc.styles[f"Heading {i}"]
        h.font.color.rgb = DARK if i > 1 else BRAND
        h.font.name = "Calibri"

    _cover_page(doc, doc_title)
    return doc


def _cover_page(doc, doc_title):
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(PROJECT)
    r.bold = True
    r.font.size = Pt(46)
    r.font.color.rgb = BRAND

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(SUBTITLE)
    r.italic = True
    r.font.size = Pt(14)
    r.font.color.rgb = GREY

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(doc_title)
    r.bold = True
    r.font.size = Pt(26)
    r.font.color.rgb = DARK

    for _ in range(6):
        doc.add_paragraph()

    meta = [
        ("Document Title", doc_title),
        ("Project", f"{PROJECT} — {SUBTITLE}"),
        ("Version", VERSION),
        ("Date", DATE_STR),
        ("Prepared By", AUTHOR),
        ("Status", "For Submission (CSO Requirement)"),
    ]
    t = doc.add_table(rows=0, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Light List Accent 1"
    for k, v in meta:
        cells = t.add_row().cells
        cells[0].text = k
        cells[1].text = v
        cells[0].paragraphs[0].runs[0].bold = True

    doc.add_page_break()


def h1(doc, text):
    doc.add_heading(text, level=1)


def h2(doc, text):
    doc.add_heading(text, level=2)


def h3(doc, text):
    doc.add_heading(text, level=3)


def para(doc, text="", bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    return p


def bullets(doc, items):
    for it in items:
        if isinstance(it, tuple):
            label, rest = it
            p = doc.add_paragraph(style="List Bullet")
            r = p.add_run(label + ": ")
            r.bold = True
            p.add_run(rest)
        else:
            doc.add_paragraph(it, style="List Bullet")


def numbered(doc, items):
    for it in items:
        doc.add_paragraph(it, style="List Number")


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(htext)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _shade(hdr[i], "404040")
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = "" if val is None else str(val)
            for pgraph in cells[i].paragraphs:
                for run in pgraph.runs:
                    run.font.size = Pt(9.5)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t


def _shade(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def toc_note(doc):
    para(doc, "Table of Contents", bold=True)
    para(doc,
         "(In Microsoft Word: References → Table of Contents → "
         "Automatic Table to generate from the headings below.)",
         italic=True)
    doc.add_page_break()


def save(doc, name):
    path = os.path.join(HERE, name)
    doc.save(path)
    print(f"  wrote {name}")
    return path


# --------------------------------------------------------------------------- #
# 01 — Project Plan
# --------------------------------------------------------------------------- #
def build_project_plan():
    doc = new_doc("Project Plan")
    toc_note(doc)

    h1(doc, "1. Introduction")
    para(doc,
         f"{PROJECT} is a {SUBTITLE.lower()} built on the MERN stack "
         "(MongoDB, Express.js, React, Node.js). It delivers Netflix-style adaptive "
         "video streaming together with secure document and presentation management, "
         "live premiere events, role-based access control, and SHA-256 file-integrity "
         "verification. This Project Plan defines the objectives, scope, schedule, "
         "resources, and risks governing the development of the platform.")

    h1(doc, "2. Project Objectives")
    bullets(doc, [
        "Provide adaptive HLS video streaming with multiple quality variants (360p, 480p, 720p, 1080p).",
        "Enable secure upload, processing, and viewing of PDF documents and PPT/PPTX presentations.",
        "Support scheduled live 'premiere' events with real-time synchronization across viewers.",
        "Enforce role-based access control (Admin / User) with JWT authentication.",
        "Guarantee content integrity using SHA-256 hash verification on every stored asset.",
        "Offer an admin dashboard with analytics and a patch/update management system.",
        "Support standard, production (Nginx), and air-gapped/offline deployments.",
    ])

    h1(doc, "3. Scope")
    h2(doc, "3.1 In Scope")
    bullets(doc, [
        "User & admin authentication, registration, and profile management.",
        "Video, document, and presentation management with background processing pipelines.",
        "Adaptive HLS streaming and in-browser document/presentation viewers.",
        "Live premieres, downloads with integrity verification, and analytics.",
        "Responsive web frontend (React + TypeScript + Tailwind CSS).",
    ])
    h2(doc, "3.2 Out of Scope")
    bullets(doc, [
        "Native mobile applications (iOS/Android).",
        "Third-party payment / subscription billing.",
        "Social networking features (comments, messaging, following).",
    ])

    h1(doc, "4. Deliverables")
    table(doc,
          ["#", "Deliverable", "Description"],
          [
              ["1", "Project Plan", "Objectives, scope, schedule, resources, risk management."],
              ["2", "SRS", "Software Requirements Specification (functional & non-functional)."],
              ["3", "Design Document", "Architecture, database schema, API, processing pipelines."],
              ["4", "Prototype", "Working web application and UI walkthrough."],
              ["5", "Use Case Document", "Actors, use-case diagram, detailed use cases."],
              ["6", "Test Cases", "Test plan and documented test cases with expected results."],
              ["7", "Source Code", "Backend (Node/Express) and frontend (React) repositories."],
          ],
          widths=[0.4, 1.8, 4.0])

    h1(doc, "5. Team & Roles")
    table(doc,
          ["Role", "Responsibility"],
          [
              ["Project Manager", "Planning, scheduling, risk and stakeholder management."],
              ["Backend Developer", "Node/Express APIs, MongoDB models, processing services."],
              ["Frontend Developer", "React/TypeScript UI, viewers, video player integration."],
              ["DevOps / Deployment", "Nginx, systemd/PM2, air-gapped & CDN deployment."],
              ["QA Engineer", "Test planning, execution, defect tracking."],
          ],
          widths=[1.8, 4.4])

    h1(doc, "6. Project Schedule & Milestones")
    para(doc, "The project is organized into seven phases over an indicative 16-week timeline.")
    table(doc,
          ["Phase", "Milestone", "Weeks", "Status"],
          [
              ["1", "Requirements & Planning (SRS, Project Plan)", "1–2", "Complete"],
              ["2", "System & Database Design", "3–4", "Complete"],
              ["3", "Authentication & RBAC", "5–6", "Complete"],
              ["4", "Media Processing & Streaming (video/PDF/PPT)", "7–10", "Complete"],
              ["5", "Premieres, Downloads, Analytics", "11–12", "Complete"],
              ["6", "Testing & Integrity Verification", "13–14", "Complete"],
              ["7", "Deployment & Documentation", "15–16", "In Progress"],
          ],
          widths=[0.7, 3.3, 1.0, 1.2])

    h1(doc, "7. Technology Stack")
    table(doc,
          ["Layer", "Technologies"],
          [
              ["Frontend", "React 19, TypeScript, Tailwind CSS, React Router, HLS.js, Socket.IO client"],
              ["Backend", "Node.js 18+, Express.js, Socket.IO, Mongoose"],
              ["Database", "MongoDB"],
              ["Auth & Security", "JWT, bcryptjs, Helmet, express-rate-limit"],
              ["Media Processing", "FFmpeg, LibreOffice, ImageMagick, Poppler-utils"],
              ["Storage", "Local filesystem or MinIO (S3-compatible)"],
              ["Infrastructure", "Nginx, systemd / PM2, Docker (optional)"],
          ],
          widths=[1.6, 4.6])

    h1(doc, "8. Risk Management")
    table(doc,
          ["ID", "Risk", "Impact", "Likelihood", "Mitigation"],
          [
              ["R1", "Media transcoding failures (FFmpeg/LibreOffice)", "High", "Medium",
               "Multi-fallback conversion pipeline; job queue with retry; status tracking."],
              ["R2", "Large file uploads exhaust storage/bandwidth", "High", "Medium",
               "Size limits, rate limiting, MinIO/object storage, CDN offload."],
              ["R3", "Unauthorized access to admin functions", "High", "Low",
               "JWT + RBAC middleware on protected routes; admin-key registration."],
              ["R4", "Content tampering / corruption", "High", "Low",
               "SHA-256 hash stored and verified on download."],
              ["R5", "Real-time sync drift during premieres", "Medium", "Medium",
               "Server-authoritative scheduler; Socket.IO time sync events."],
              ["R6", "Deployment in air-gapped environment", "Medium", "Medium",
               "Offline installation packages and documented commands."],
          ],
          widths=[0.4, 1.9, 0.7, 0.8, 2.4])

    h1(doc, "9. Tools Used")
    bullets(doc, [
        "Version Control: Git",
        "IDE: Visual Studio Code",
        "Package Managers: npm",
        "Process Management: nodemon (dev), PM2 / systemd (prod)",
        "Testing: Manual test cases, REST clients, browser dev tools",
        "Documentation: Microsoft Word / Markdown",
    ])

    return save(doc, "01_Project_Plan.docx")


# --------------------------------------------------------------------------- #
# 02 — SRS
# --------------------------------------------------------------------------- #
def build_srs():
    doc = new_doc("Software Requirements Specification (SRS)")
    toc_note(doc)

    h1(doc, "1. Introduction")
    h2(doc, "1.1 Purpose")
    para(doc,
         f"This Software Requirements Specification (SRS) describes the functional and "
         f"non-functional requirements of {PROJECT}, a secure media streaming and document "
         "management platform. It is intended for the development team, testers, and "
         "evaluators as the authoritative reference for what the system must do.")
    h2(doc, "1.2 Scope")
    para(doc,
         f"{PROJECT} allows administrators to upload and manage video, document, and "
         "presentation content, and allows authenticated users to stream and view that "
         "content securely. The system provides adaptive HLS streaming, live premieres, "
         "downloads with integrity verification, analytics, and role-based access control.")
    h2(doc, "1.3 Definitions, Acronyms & Abbreviations")
    table(doc,
          ["Term", "Definition"],
          [
              ["HLS", "HTTP Live Streaming – adaptive bitrate streaming protocol."],
              ["JWT", "JSON Web Token – token used for stateless authentication."],
              ["RBAC", "Role-Based Access Control."],
              ["SHA-256", "Secure Hash Algorithm used for file-integrity verification."],
              ["MERN", "MongoDB, Express.js, React, Node.js technology stack."],
              ["Premiere", "A scheduled, synchronized live video event."],
              ["MinIO", "S3-compatible object storage service."],
          ],
          widths=[1.3, 4.9])
    h2(doc, "1.4 References")
    bullets(doc, [
        "IEEE Std 830-1998, Recommended Practice for Software Requirements Specifications.",
        "PakStream README.md and PROJECT_ARCHITECTURE.md.",
    ])

    h1(doc, "2. Overall Description")
    h2(doc, "2.1 Product Perspective")
    para(doc,
         "PakStream is a self-contained, three-tier web application: a React frontend, "
         "a Node.js/Express REST + Socket.IO backend, and a MongoDB database, fronted by "
         "Nginx for static delivery and HLS streaming. Media is processed by FFmpeg, "
         "LibreOffice, ImageMagick, and Poppler-utils.")
    h2(doc, "2.2 Product Functions")
    bullets(doc, [
        "Authentication, registration, and profile management.",
        "Video upload, transcoding to HLS, and adaptive playback.",
        "Document (PDF) and presentation (PPT/PPTX) upload, conversion, and viewing.",
        "Scheduled live premieres with real-time synchronization.",
        "Downloads with SHA-256 integrity verification.",
        "Admin dashboard, analytics, user management, and patch management.",
    ])
    h2(doc, "2.3 User Classes & Characteristics")
    table(doc,
          ["User Class", "Characteristics"],
          [
              ["Guest", "Unauthenticated visitor; may register or log in only."],
              ["User", "Authenticated viewer; streams/views content, joins premieres, downloads, manages own profile."],
              ["Admin", "Privileged user; uploads/manages content, manages users, views analytics, applies patches."],
          ],
          widths=[1.3, 4.9])
    h2(doc, "2.4 Operating Environment")
    bullets(doc, [
        "Server: Linux (Ubuntu) with Node.js 18+, MongoDB, Nginx, FFmpeg, LibreOffice.",
        "Client: Modern web browser (Chrome, Firefox, Edge, Safari) with JavaScript enabled.",
    ])
    h2(doc, "2.5 Assumptions & Dependencies")
    bullets(doc, [
        "External processing tools (FFmpeg, LibreOffice, ImageMagick, Poppler) are installed on the server.",
        "MongoDB is reachable by the backend.",
        "Sufficient storage exists for original and transcoded media.",
    ])

    h1(doc, "3. Functional Requirements")

    def fr_block(title, rows):
        h2(doc, title)
        table(doc, ["ID", "Requirement", "Priority"], rows, widths=[0.7, 4.7, 0.8])

    fr_block("3.1 Authentication & Access Control", [
        ["FR-1", "The system shall allow a guest to register as a user with email and password.", "High"],
        ["FR-2", "The system shall allow admin registration only with a valid admin key.", "High"],
        ["FR-3", "The system shall authenticate users via JWT and expire tokens after 24 hours.", "High"],
        ["FR-4", "The system shall hash passwords with bcrypt before storage.", "High"],
        ["FR-5", "The system shall restrict admin-only endpoints to users with the Admin role.", "High"],
        ["FR-6", "The system shall allow users to view and edit their profile (avatar, bio).", "Medium"],
    ])
    fr_block("3.2 Video Management", [
        ["FR-7", "The system shall allow an admin to upload a video with progress tracking.", "High"],
        ["FR-8", "The system shall transcode uploaded video to HLS in 360p/480p/720p/1080p.", "High"],
        ["FR-9", "The system shall generate a thumbnail and extract metadata (duration, resolution).", "Medium"],
        ["FR-10", "The system shall track video status: uploading → processing → ready.", "High"],
        ["FR-11", "The system shall allow users to stream videos adaptively via HLS.js.", "High"],
        ["FR-12", "The system shall record views, likes, and dislikes per video.", "Low"],
    ])
    fr_block("3.3 Document & Presentation Management", [
        ["FR-13", "The system shall allow upload of PDF documents and extract page count.", "High"],
        ["FR-14", "The system shall generate a thumbnail from the first PDF page.", "Medium"],
        ["FR-15", "The system shall allow upload of PPT/PPTX and convert slides to PNG images.", "High"],
        ["FR-16", "The system shall provide in-browser viewers for documents and presentations.", "High"],
        ["FR-17", "The system shall record views and likes for documents and presentations.", "Low"],
    ])
    fr_block("3.4 Premieres, Downloads & Analytics", [
        ["FR-18", "The system shall allow an admin to schedule a live premiere for a video.", "Medium"],
        ["FR-19", "The system shall synchronize premiere playback across viewers in real time.", "Medium"],
        ["FR-20", "The system shall compute and store a SHA-256 hash for each stored asset.", "High"],
        ["FR-21", "The system shall verify the SHA-256 hash on download and record download history.", "High"],
        ["FR-22", "The system shall provide an admin analytics dashboard of usage metrics.", "Medium"],
        ["FR-23", "The system shall provide a patch/update management capability for admins.", "Low"],
    ])

    h1(doc, "4. Non-Functional Requirements")
    h2(doc, "4.1 Security")
    bullets(doc, [
        "JWT-based stateless authentication and bcrypt password hashing.",
        "Helmet security headers and express-rate-limit throttling.",
        "SHA-256 integrity verification of all stored media.",
        "Role-based authorization on protected routes.",
    ])
    h2(doc, "4.2 Performance")
    bullets(doc, [
        "Adaptive HLS streaming to match client bandwidth.",
        "Background job queue for non-blocking media transcoding.",
        "Nginx static serving and optional CDN offload for media delivery.",
    ])
    h2(doc, "4.3 Scalability")
    bullets(doc, [
        "Stateless API enabling horizontal scaling behind a load balancer.",
        "Pluggable storage abstraction (local filesystem or MinIO object storage).",
    ])
    h2(doc, "4.4 Usability")
    bullets(doc, [
        "Responsive React UI with Tailwind CSS across desktop and mobile browsers.",
        "Real-time processing/status feedback via Socket.IO.",
    ])
    h2(doc, "4.5 Reliability & Availability")
    bullets(doc, [
        "MongoDB connection retry logic on the backend.",
        "Health-check endpoint (GET /api/health) for monitoring.",
        "Support for air-gapped/offline deployment for restricted environments.",
    ])

    h1(doc, "5. External Interface Requirements")
    h2(doc, "5.1 User Interfaces")
    para(doc, "Web-based responsive interface served by the React frontend.")
    h2(doc, "5.2 Software Interfaces")
    bullets(doc, [
        "REST API over HTTP(S) (JSON) for all CRUD operations.",
        "WebSocket interface via Socket.IO for real-time events.",
        "MongoDB via the Mongoose ODM.",
        "FFmpeg / LibreOffice / ImageMagick / Poppler invoked as processing services.",
    ])

    return save(doc, "02_SRS.docx")


# --------------------------------------------------------------------------- #
# 03 — Design Document
# --------------------------------------------------------------------------- #
def build_design():
    doc = new_doc("Design Document")
    toc_note(doc)

    h1(doc, "1. Introduction")
    para(doc,
         f"This document describes the architecture and detailed design of {PROJECT}, "
         "covering the system architecture, module decomposition, database schema, API "
         "design, media-processing pipelines, real-time communication, and security design.")

    h1(doc, "2. System Architecture")
    para(doc,
         "PakStream follows a three-tier architecture fronted by an Nginx reverse proxy. "
         "The presentation tier is a React single-page application; the application tier is a "
         "Node.js/Express server exposing REST APIs and a Socket.IO real-time channel; the "
         "data tier is MongoDB plus a media storage layer (local filesystem or MinIO).")
    para(doc, "Logical architecture (top-down):", bold=True)
    para(doc,
         "[ Browser / React SPA ]\n"
         "        |  HTTPS (REST + WebSocket)\n"
         "        v\n"
         "[ Nginx reverse proxy / static + HLS serving ]\n"
         "        |\n"
         "        v\n"
         "[ Node.js + Express API  |  Socket.IO  |  Premiere Scheduler ]\n"
         "        |                         |\n"
         "        v                         v\n"
         "[ MongoDB (Mongoose) ]   [ Media Processing: FFmpeg / LibreOffice / ImageMagick / Poppler ]\n"
         "                                  |\n"
         "                                  v\n"
         "                    [ Storage: Local FS or MinIO (S3) ]")

    h1(doc, "3. Module Decomposition")
    h2(doc, "3.1 Backend Modules")
    table(doc,
          ["Module", "Responsibility"],
          [
              ["Controllers", "Request handlers for auth, video, document, presentation, user, premiere, download, analytics, patch."],
              ["Routes", "Express route definitions mapping URLs to controllers."],
              ["Middleware", "Authentication, role checks, and file-upload (multer) handling."],
              ["Models", "Mongoose schemas: User, Video, Document, Presentation, Premiere, Download, VideoDownload, Patch."],
              ["Services", "videoProcessor, documentProcessor, presentationProcessor, hashService, storageService, videoQueue, premiereScheduler."],
              ["Socket", "Socket.IO event handlers for real-time updates and premiere sync."],
          ],
          widths=[1.5, 4.7])
    h2(doc, "3.2 Frontend Modules")
    table(doc,
          ["Module", "Responsibility"],
          [
              ["Pages", "Admin and user page groups (dashboards, viewers, management)."],
              ["Components", "Reusable UI: admin, auth, common, document, patch, premiere, presentation, video."],
              ["Services", "Typed API clients for each backend resource."],
              ["Contexts", "Auth, theme, and notification providers (React Context API)."],
              ["Hooks", "Custom hooks for data fetching and real-time updates."],
          ],
          widths=[1.5, 4.7])

    h1(doc, "4. Database Design")
    para(doc, "Core MongoDB collections (via Mongoose) and their key fields:")
    table(doc,
          ["Model", "Key Fields", "Purpose"],
          [
              ["User", "name, email, passwordHash, role, avatar, bio", "Accounts & RBAC."],
              ["Video", "title, description, category, status, hlsPath, thumbnail, duration, resolution, sha256, views, likes", "Streaming video assets."],
              ["Document", "title, category, filePath, pageCount, thumbnail, sha256, views, likes", "PDF documents."],
              ["Presentation", "title, category, slides[], thumbnail, duration, sha256, views, likes", "PPT/PPTX presentations."],
              ["Premiere", "video, scheduledTime, status, startedAt", "Scheduled live events."],
              ["Download", "user, resourceType, resourceId, hashVerified, timestamp", "Document/asset download history."],
              ["VideoDownload", "user, video, hashVerified, timestamp", "Original-video download history."],
              ["Patch", "name, version, appliedAt, status", "Update/patch management."],
          ],
          widths=[1.2, 3.2, 1.8])

    h1(doc, "5. API Design")
    para(doc, "RESTful endpoints are grouped by resource under the /api prefix:")
    table(doc,
          ["Base Route", "Responsibility"],
          [
              ["/api/auth", "Register, login, profile."],
              ["/api/videos", "Video CRUD, upload, playback, likes."],
              ["/api/documents", "Document CRUD, upload, view."],
              ["/api/presentations", "Presentation CRUD, upload, view."],
              ["/api/users", "User management (admin)."],
              ["/api/premieres", "Schedule and manage live premieres."],
              ["/api/downloads", "Download tracking & integrity verification."],
              ["/api/analytics", "Usage analytics for the admin dashboard."],
              ["/api/patch", "Patch / update management."],
              ["/api/health", "Health check for monitoring."],
          ],
          widths=[2.0, 4.2])

    h1(doc, "6. Media Processing Pipelines")
    h2(doc, "6.1 Video")
    numbered(doc, [
        "Admin uploads original video (multer); status set to 'uploading'.",
        "Job enqueued in the video queue; status 'processing'.",
        "FFmpeg transcodes to HLS variants (360p/480p/720p/1080p) and generates a thumbnail.",
        "Metadata (duration, resolution) extracted; SHA-256 computed and stored.",
        "Status set to 'ready'; client notified via Socket.IO.",
    ])
    h2(doc, "6.2 Document (PDF)")
    numbered(doc, [
        "PDF uploaded; page count extracted via Poppler-utils.",
        "First-page thumbnail generated; SHA-256 computed.",
        "Document made available in the in-browser viewer.",
    ])
    h2(doc, "6.3 Presentation (PPT/PPTX)")
    numbered(doc, [
        "Presentation uploaded.",
        "Converted to PNG slides using a multi-fallback chain (ImageMagick → Poppler → LibreOffice).",
        "Thumbnail and slide list generated; SHA-256 computed.",
        "Presentation made available for slide-by-slide viewing.",
    ])

    h1(doc, "7. Real-Time Design (Socket.IO)")
    bullets(doc, [
        "Processing status updates pushed to clients as media moves through the pipeline.",
        "Premiere synchronization: server-authoritative scheduler broadcasts start/seek events.",
        "Connection lifecycle managed per authenticated user session.",
    ])

    h1(doc, "8. Security Design")
    bullets(doc, [
        "JWT authentication with 24-hour expiry; bcrypt password hashing.",
        "RBAC middleware guarding admin-only routes; admin-key gated admin registration.",
        "Helmet security headers and express-rate-limit on the API.",
        "SHA-256 integrity verification on stored and downloaded assets.",
    ])

    return save(doc, "03_Design_Document.docx")


# --------------------------------------------------------------------------- #
# 04 — Prototype Document
# --------------------------------------------------------------------------- #
def build_prototype():
    doc = new_doc("Prototype Document")
    toc_note(doc)

    h1(doc, "1. Introduction")
    para(doc,
         f"This document presents the working prototype of {PROJECT}. The prototype is a "
         "fully functional web application built with React (TypeScript) on the frontend and "
         "Node.js/Express on the backend. This document walks through the existing user "
         "interface screen by screen. Representative screenshots are available in the "
         "project's demo/ folder.")

    h1(doc, "2. Available Screenshots")
    para(doc, "The following screenshots are included in the repository's demo/ folder:")
    bullets(doc, [
        ("demo/dashboard.png", "Main dashboard / content browse view."),
        ("demo/documents.png", "Document library and viewer."),
        ("demo/logo.png", "Application branding / logo."),
    ])
    # Embed screenshots if present
    for fname, caption in [("dashboard.png", "Figure 1 – Dashboard"),
                           ("documents.png", "Figure 2 – Documents")]:
        fpath = os.path.join(DEMO_DIR, fname)
        if os.path.exists(fpath):
            try:
                doc.add_picture(fpath, width=Inches(6.0))
                cap = doc.add_paragraph(caption)
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap.runs[0].italic = True
                cap.runs[0].font.size = Pt(9)
            except Exception as exc:  # pragma: no cover
                para(doc, f"[Could not embed {fname}: {exc}]", italic=True)

    h1(doc, "3. Screen Walkthrough")

    screens = [
        ("3.1 Login & Registration",
         "Entry screens for authentication. Users register with email and password; admins "
         "register using an admin key. On success a JWT is issued and the user is routed to "
         "their dashboard. Form validation and error messaging are provided."),
        ("3.2 User Dashboard",
         "The landing experience after login. Presents featured and categorized content "
         "(videos, documents, presentations) in a Netflix-style browse layout with thumbnails, "
         "search, and category filters."),
        ("3.3 Video Player",
         "Adaptive HLS player (HLS.js) with quality switching across 360p/480p/720p/1080p, "
         "play/pause/seek controls, view counting, and like/dislike actions."),
        ("3.4 Document Viewer",
         "In-browser PDF viewer showing page-by-page navigation, page count, and a download "
         "option with SHA-256 integrity verification."),
        ("3.5 Presentation Viewer",
         "Slide-by-slide viewer rendering converted PNG slides with next/previous navigation "
         "and thumbnails."),
        ("3.6 Premieres",
         "List of scheduled live premiere events with countdowns; on start, playback is "
         "synchronized across viewers in real time via Socket.IO."),
        ("3.7 Admin Dashboard",
         "Admin-only control center for uploading and managing videos, documents, and "
         "presentations, with real-time processing status indicators."),
        ("3.8 Analytics",
         "Charts and metrics (via Recharts) summarizing views, downloads, and content usage "
         "for administrators."),
        ("3.9 User Management",
         "Admin interface to list, search, and manage user accounts and roles."),
        ("3.10 Profile",
         "User profile screen for editing avatar, bio, and account details."),
    ]
    for title, desc in screens:
        h2(doc, title)
        para(doc, desc)

    h1(doc, "4. Design & Interaction Notes")
    bullets(doc, [
        "Consistent dark, media-centric theme styled with Tailwind CSS.",
        "Responsive layouts adapting from desktop to mobile browsers.",
        "Real-time feedback for long-running uploads and transcoding jobs.",
        "Role-aware navigation: admin tools are hidden from standard users.",
    ])

    return save(doc, "04_Prototype_Document.docx")


# --------------------------------------------------------------------------- #
# 05 — Use Case Document
# --------------------------------------------------------------------------- #
def build_use_cases():
    doc = new_doc("Use Case Document")
    toc_note(doc)

    h1(doc, "1. Introduction")
    para(doc,
         f"This document defines the actors and use cases of {PROJECT}. It includes a "
         "use-case summary, a textual use-case diagram, and detailed descriptions of the "
         "primary use cases.")

    h1(doc, "2. Actors")
    table(doc,
          ["Actor", "Description"],
          [
              ["Guest", "Unauthenticated visitor; can register or log in."],
              ["User", "Authenticated viewer; consumes content, joins premieres, downloads, manages profile."],
              ["Admin", "Privileged user; manages content and users, schedules premieres, views analytics."],
              ["System", "Background services: media processing, scheduler, integrity verification."],
          ],
          widths=[1.2, 5.0])

    h1(doc, "3. Use Case Diagram (Textual)")
    para(doc,
         "Guest\n"
         "  └─ Register / Login\n\n"
         "User\n"
         "  ├─ Browse Content\n"
         "  ├─ Watch Video (HLS)\n"
         "  ├─ View Document\n"
         "  ├─ View Presentation\n"
         "  ├─ Join Premiere\n"
         "  ├─ Download (with integrity verification)\n"
         "  └─ Manage Profile\n\n"
         "Admin (all User capabilities, plus)\n"
         "  ├─ Upload & Manage Video / Document / Presentation\n"
         "  ├─ Schedule Premiere\n"
         "  ├─ Manage Users\n"
         "  ├─ View Analytics\n"
         "  └─ Manage Patches\n\n"
         "System\n"
         "  ├─ Process Media (FFmpeg/LibreOffice/...)\n"
         "  ├─ Verify Integrity (SHA-256)\n"
         "  └─ Synchronize Premiere")

    h1(doc, "4. Use Case Summary")
    table(doc,
          ["ID", "Use Case", "Primary Actor"],
          [
              ["UC-1", "Register / Login", "Guest"],
              ["UC-2", "Upload Video", "Admin"],
              ["UC-3", "Watch Video (HLS)", "User"],
              ["UC-4", "View Document", "User"],
              ["UC-5", "View Presentation", "User"],
              ["UC-6", "Join Premiere", "User"],
              ["UC-7", "Download with Integrity Verification", "User"],
              ["UC-8", "Manage Users", "Admin"],
              ["UC-9", "View Analytics", "Admin"],
          ],
          widths=[0.7, 3.5, 1.5])

    h1(doc, "5. Detailed Use Cases")

    def use_case(uc_id, name, actor, pre, main, alt, post):
        h2(doc, f"{uc_id}: {name}")
        table(doc, ["Field", "Value"], [
            ["Use Case ID", uc_id],
            ["Primary Actor", actor],
            ["Preconditions", pre],
            ["Postconditions", post],
        ], widths=[1.4, 4.8])
        para(doc, "Main Flow:", bold=True)
        numbered(doc, main)
        if alt:
            para(doc, "Alternative / Exception Flows:", bold=True)
            bullets(doc, alt)

    use_case(
        "UC-1", "Register / Login", "Guest",
        "User has a valid email; (for admin) a valid admin key.",
        ["Guest opens the registration or login screen.",
         "Guest submits credentials (and admin key, if registering as admin).",
         "System validates input and, for registration, hashes the password with bcrypt.",
         "System issues a JWT and redirects to the appropriate dashboard."],
        ["Invalid credentials → system shows an authentication error.",
         "Invalid/missing admin key → admin registration is rejected.",
         "Duplicate email → registration is rejected."],
        "User is authenticated with a valid 24-hour JWT.")

    use_case(
        "UC-2", "Upload Video", "Admin",
        "Admin is logged in with the Admin role.",
        ["Admin selects a video file and metadata (title, category).",
         "System uploads the file with progress tracking and sets status to 'uploading'.",
         "System enqueues a transcoding job (status 'processing').",
         "FFmpeg produces HLS variants, a thumbnail, and metadata; SHA-256 is computed.",
         "Status is set to 'ready' and the admin is notified in real time."],
        ["Transcoding fails → status flagged as failed; admin is notified.",
         "Unsupported format → upload rejected with an error."],
        "Video is available for adaptive HLS streaming.")

    use_case(
        "UC-3", "Watch Video (HLS)", "User",
        "User is logged in; the video status is 'ready'.",
        ["User selects a video from the dashboard.",
         "Frontend loads the HLS manifest via HLS.js.",
         "Player streams adaptively based on bandwidth; controls allow play/pause/seek.",
         "System increments the view counter."],
        ["Network degradation → player drops to a lower quality variant automatically."],
        "User views the video; view count is updated.")

    use_case(
        "UC-6", "Join Premiere", "User",
        "User is logged in; a premiere is scheduled or live.",
        ["User opens the premieres list and selects an event.",
         "Before start time, a countdown is shown.",
         "At start, the server-authoritative scheduler broadcasts a start event via Socket.IO.",
         "Playback is synchronized across all connected viewers."],
        ["User joins late → client seeks to the synchronized server position."],
        "User watches the premiere in sync with other viewers.")

    use_case(
        "UC-7", "Download with Integrity Verification", "User",
        "User is logged in; the asset exists with a stored SHA-256 hash.",
        ["User requests a download of a video, document, or presentation.",
         "System streams the original file to the user.",
         "System verifies the SHA-256 hash and records the download in history."],
        ["Hash mismatch → download flagged as failed integrity check."],
        "User receives an integrity-verified file; download is logged.")

    use_case(
        "UC-8", "Manage Users", "Admin",
        "Admin is logged in with the Admin role.",
        ["Admin opens the user-management screen.",
         "Admin searches/filters the user list.",
         "Admin updates a user's role or account state.",
         "System persists the change."],
        ["Admin attempts to remove the last admin → action is prevented."],
        "User accounts reflect the administrative changes.")

    return save(doc, "05_Use_Case_Document.docx")


# --------------------------------------------------------------------------- #
# 06 — Test Cases
# --------------------------------------------------------------------------- #
def build_test_cases():
    doc = new_doc("Test Cases Document")
    toc_note(doc)

    h1(doc, "1. Introduction")
    para(doc,
         f"This document defines the test plan and detailed test cases for {PROJECT}. "
         "Test cases cover authentication and access control, media upload and processing, "
         "streaming and viewing, premieres, downloads with integrity verification, and "
         "negative/security scenarios.")

    h1(doc, "2. Test Approach")
    bullets(doc, [
        "Functional testing against the documented requirements (SRS).",
        "Positive and negative test cases for each major module.",
        "Security-focused tests for authentication, RBAC, and integrity verification.",
        "Manual execution via the web UI and REST clients; results recorded as Pass/Fail.",
    ])

    h1(doc, "3. Test Cases")
    rows = [
        ["TC-01", "Auth", "Valid user registration", "On register screen",
         "Submit valid name/email/password", "newuser@test.com", "Account created; JWT issued; redirected to dashboard", "Positive"],
        ["TC-02", "Auth", "Duplicate email registration", "Email already exists",
         "Submit existing email", "existing@test.com", "Registration rejected with duplicate-email error", "Negative"],
        ["TC-03", "Auth", "Admin registration with valid key", "Valid admin key available",
         "Register with admin key", "valid admin key", "Admin account created with Admin role", "Positive"],
        ["TC-04", "Auth", "Admin registration with invalid key", "On admin register",
         "Register with wrong admin key", "wrong key", "Registration rejected", "Negative"],
        ["TC-05", "Auth", "Valid login", "Account exists",
         "Submit correct credentials", "valid creds", "Login succeeds; JWT issued", "Positive"],
        ["TC-06", "Auth", "Invalid login", "Account exists",
         "Submit wrong password", "wrong password", "Login rejected with auth error", "Negative"],
        ["TC-07", "Auth", "Token expiry", "Logged in",
         "Use API after 24h with stale token", "expired JWT", "Request rejected; re-authentication required", "Negative"],
        ["TC-08", "RBAC", "User blocked from admin endpoint", "Logged in as User",
         "Call admin-only endpoint", "user JWT", "403 Forbidden returned", "Security"],
        ["TC-09", "RBAC", "Admin accesses admin endpoint", "Logged in as Admin",
         "Call admin endpoint", "admin JWT", "Request succeeds", "Positive"],
        ["TC-10", "Video", "Upload valid video", "Logged in as Admin",
         "Upload MP4 with metadata", "sample.mp4", "Status progresses uploading→processing→ready", "Positive"],
        ["TC-11", "Video", "Upload unsupported format", "Logged in as Admin",
         "Upload .exe file", "bad.exe", "Upload rejected with format error", "Negative"],
        ["TC-12", "Video", "HLS variants generated", "Video processed",
         "Inspect output", "ready video", "360p/480p/720p/1080p HLS variants exist", "Positive"],
        ["TC-13", "Video", "Adaptive playback", "Video ready",
         "Play video, throttle bandwidth", "ready video", "Player drops to lower quality automatically", "Positive"],
        ["TC-14", "Video", "View counter increments", "Video ready",
         "Open video", "ready video", "View count increases by 1", "Positive"],
        ["TC-15", "Document", "Upload PDF", "Logged in as Admin",
         "Upload PDF", "sample.pdf", "Page count extracted; thumbnail generated", "Positive"],
        ["TC-16", "Document", "View PDF", "Document ready",
         "Open document viewer", "ready doc", "Pages render; navigation works", "Positive"],
        ["TC-17", "Presentation", "Upload PPTX", "Logged in as Admin",
         "Upload PPTX", "deck.pptx", "Slides converted to PNG images", "Positive"],
        ["TC-18", "Presentation", "View slides", "Presentation ready",
         "Open presentation viewer", "ready deck", "Slides render with next/prev navigation", "Positive"],
        ["TC-19", "Premiere", "Schedule premiere", "Logged in as Admin",
         "Schedule premiere for a video", "future time", "Premiere created with countdown", "Positive"],
        ["TC-20", "Premiere", "Synchronized start", "Premiere scheduled",
         "Two clients join at start time", "2 viewers", "Playback synchronized across viewers", "Positive"],
        ["TC-21", "Premiere", "Late join sync", "Premiere live",
         "Join after start", "1 late viewer", "Client seeks to server position", "Positive"],
        ["TC-22", "Download", "Download with valid integrity", "Asset has stored hash",
         "Download asset", "ready asset", "File downloads; SHA-256 verified; logged", "Positive"],
        ["TC-23", "Download", "Integrity mismatch detected", "Asset corrupted",
         "Download tampered asset", "tampered file", "Integrity check fails; flagged", "Security"],
        ["TC-24", "Analytics", "Admin views analytics", "Logged in as Admin",
         "Open analytics dashboard", "admin JWT", "Charts/metrics render correctly", "Positive"],
        ["TC-25", "Users", "Admin manages users", "Logged in as Admin",
         "Change a user's role", "target user", "Role updated and persisted", "Positive"],
        ["TC-26", "API", "Health check", "Server running",
         "GET /api/health", "–", "200 OK with healthy status", "Positive"],
        ["TC-27", "Security", "Rate limiting", "Server running",
         "Exceed request rate", "burst requests", "Excess requests throttled (429)", "Security"],
        ["TC-28", "Profile", "Edit profile", "Logged in as User",
         "Update avatar/bio", "new bio", "Profile updated and persisted", "Positive"],
    ]
    table(doc,
          ["ID", "Module", "Description", "Preconditions", "Steps", "Test Data", "Expected Result", "Type"],
          rows,
          widths=[0.45, 0.7, 1.1, 0.9, 1.1, 0.7, 1.5, 0.6])

    h1(doc, "4. Summary")
    para(doc,
         f"A total of {len(rows)} test cases are defined across authentication, RBAC, video, "
         "document, presentation, premiere, download, analytics, user management, API, "
         "security, and profile modules. Each case records preconditions, steps, test data, "
         "and expected results, and is executed with a Pass/Fail verdict.")

    return save(doc, "06_Test_Cases.docx")


# --------------------------------------------------------------------------- #
def main():
    print("Generating PakStream CSO documentation (.docx) ...")
    build_project_plan()
    build_srs()
    build_design()
    build_prototype()
    build_use_cases()
    build_test_cases()
    print("Done. Files written to:", HERE)


if __name__ == "__main__":
    main()
